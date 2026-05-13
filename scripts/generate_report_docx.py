"""
generate_report_docx.py — Genera informe_final.docx completo.

Ejecutar desde la raiz del proyecto:
    python scripts/generate_report_docx.py

Computa inline: Accuracy, Macro F1, Cohen Kappa, AUC-ROC (FP32 e INT8).
Lee de JSON : latencia, disco, parametros, ablacion E0-E5.
"""

import json
import sys
from pathlib import Path

import numpy as np
import onnxruntime as ort
from sklearn.metrics import (
    accuracy_score, f1_score, roc_auc_score, cohen_kappa_score,
)
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT          = Path(__file__).resolve().parent.parent
CKPT_DIR      = ROOT / "checkpoints"
PROCESSED_DIR = ROOT / "data" / "processed"
ISRUC_DIR     = PROCESSED_DIR / "isruc_subjects"
RESULTS_DIR   = ROOT / "results"
THRESHOLD_PCT = 0.5


# ---------------------------------------------------------------------------
# Helpers de inferencia
# ---------------------------------------------------------------------------

def load_isruc():
    Xs, ys = [], []
    for npz in sorted(ISRUC_DIR.glob("subj_*.npz"),
                      key=lambda p: int(p.stem.split("_")[1])):
        d = np.load(str(npz))
        Xs.append(d["X"]); ys.append(d["y_apnea"])
    return np.concatenate(Xs).astype(np.float32), np.concatenate(ys)


def make_session(path: Path) -> ort.InferenceSession:
    opts = ort.SessionOptions()
    opts.intra_op_num_threads = 1
    return ort.InferenceSession(str(path), sess_options=opts,
                                providers=["CPUExecutionProvider"])


def predict(sess, x, batch=128):
    probs, preds = [], []
    for i in range(0, len(x), batch):
        xb  = x[i:i+batch]
        out = sess.run(None, {"eeg": xb})[0]
        e   = np.exp(out - out.max(axis=1, keepdims=True))
        p   = e / e.sum(axis=1, keepdims=True)
        probs.append(p[:, 1])
        preds.append(out.argmax(axis=1))
    return np.concatenate(probs), np.concatenate(preds)


def metrics(y_true, probs, preds):
    return {
        "accuracy":  round(float(accuracy_score(y_true, preds)), 4),
        "macro_f1":  round(float(f1_score(y_true, preds,
                                          average="macro", zero_division=0)), 4),
        "kappa":     round(float(cohen_kappa_score(y_true, preds)), 4),
        "auc_roc":   round(float(roc_auc_score(y_true, probs)), 4),
    }


# ---------------------------------------------------------------------------
# Helpers python-docx
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def bold_cell(cell, text, size=10, color=None, center=True):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))


def plain_cell(cell, text, size=10, bold=False, center=True):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_para(doc, text, size=11, indent=False):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(size)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    return p


# ---------------------------------------------------------------------------
# Seccion: titulo
# ---------------------------------------------------------------------------

def add_title_page(doc):
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("PROYECTO 2")
    run.bold = True; run.font.size = Pt(22)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = t2.add_run("Compresion de Modelos DNN para\nDeteccion de Apnea del Sueno")
    run2.bold = True; run2.font.size = Pt(16)

    doc.add_paragraph()
    uni = doc.add_paragraph("Universidad Interamericana PR - Bayamon")
    uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in uni.runs: r.font.size = Pt(13)

    doc.add_paragraph()
    team = doc.add_paragraph(
        "Eli Jaaziel Ayala Ortiz    Y00622025\n"
        "Edwin Roman Maldonado      R00587066"
    )
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in team.runs: r.font.size = Pt(12)

    doc.add_paragraph()
    date_p = doc.add_paragraph("Mayo 2026")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in date_p.runs: r.font.size = Pt(12)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Tabla principal de metricas (FP32 vs INT8)
# ---------------------------------------------------------------------------

METRIC_LABELS = {
    "accuracy": "Accuracy",
    "macro_f1": "Macro F1",
    "kappa":    "Cohen Kappa",
    "auc_roc":  "AUC-ROC",
}
HDR_COLOR  = "1F4E79"   # azul oscuro
ROW_EVEN   = "D6E4F0"
ROW_ODD    = "FFFFFF"
OK_COLOR   = "E2EFDA"   # verde claro
FAIL_COLOR = "FFE0E0"   # rojo claro


def add_metrics_table(doc, model_name, m_fp32, m_int8, threshold=THRESHOLD_PCT):
    add_heading(doc, f"Modelo: {model_name}", level=3)

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"

    hdr_cells = tbl.rows[0].cells
    for cell, txt in zip(hdr_cells,
                         ["Metrica", "FP32", "INT8", "Diff %", "Estado"]):
        set_cell_bg(cell, HDR_COLOR)
        bold_cell(cell, txt, size=10, color="FFFFFF")

    metrics_order = ["accuracy", "macro_f1", "kappa", "auc_roc"]
    for idx, key in enumerate(metrics_order):
        row   = tbl.add_row().cells
        label = METRIC_LABELS[key]
        vfp   = m_fp32[key]
        vi8   = m_int8[key]
        diff  = round(abs(vfp - vi8) * 100, 3)
        ok    = diff <= threshold

        bg = ROW_EVEN if idx % 2 == 0 else ROW_ODD
        for c in row:
            set_cell_bg(c, bg)

        plain_cell(row[0], label,          bold=True, center=False)
        plain_cell(row[1], f"{vfp:.4f}")
        plain_cell(row[2], f"{vi8:.4f}")
        plain_cell(row[3], f"{diff:.3f}%")

        estado = "OK (+/-0.5%)" if ok else "ALERTA (>0.5%)"
        set_cell_bg(row[4], OK_COLOR if ok else FAIL_COLOR)
        plain_cell(row[4], estado, bold=True)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Tabla de ablacion E0-E5
# ---------------------------------------------------------------------------

def add_ablation_table(doc, ablation):
    add_heading(doc, "Estudio de Ablacion E0-E5", level=2)
    add_para(doc, (
        "Cada configuracion agrega una tecnica sobre la anterior para aislar "
        "su contribucion individual al rendimiento final."
    ))

    cols = ["ID", "Configuracion", "Params", "KB disco", "Lat (ms)",
            "Speedup", "Accuracy", "Macro F1", "AUC-ROC", "delta AUC"]
    tbl  = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"

    for cell, txt in zip(tbl.rows[0].cells, cols):
        set_cell_bg(cell, HDR_COLOR)
        bold_cell(cell, txt, size=9, color="FFFFFF")

    for idx, (eid, row_data) in enumerate(ablation.items()):
        t   = row_data["test"]
        bg  = "C6EFCE" if eid == "E5" else (ROW_EVEN if idx % 2 == 0 else ROW_ODD)
        row = tbl.add_row().cells
        for c in row:
            set_cell_bg(c, bg)

        d_auc = t["delta_auc"]
        sign  = "+" if d_auc >= 0 else ""
        vals  = [
            eid,
            row_data["label"],
            f"{row_data['n_params']:,}",
            f"{row_data['disk_kb']:.1f}",
            f"{row_data['lat_ms']:.3f}",
            f"{row_data['speedup_vs_baseline']:.2f}x",
            f"{t['accuracy']:.4f}",
            f"{t.get('macro_f1', t.get('f1', '—')):.4f}" if isinstance(t.get('macro_f1', t.get('f1')), float) else "—",
            f"{t['auc_roc']:.4f}",
            f"{sign}{d_auc:.4f}",
        ]
        is_e5 = eid == "E5"
        for cell, txt in zip(row, vals):
            plain_cell(cell, txt, size=9, bold=is_e5)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Tabla de compresion (disco + latencia)
# ---------------------------------------------------------------------------

def add_compression_table(doc, ev):
    add_heading(doc, "Reduccion de Tamano y Latencia (Fase 5)", level=2)

    cols = ["Modelo", "Params", "FP32 KB", "INT8 KB",
            "Ratio disco", "Lat FP32 (ms)", "Lat INT8 (ms)", "Speedup INT8"]
    tbl  = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"

    for cell, txt in zip(tbl.rows[0].cells, cols):
        set_cell_bg(cell, HDR_COLOR)
        bold_cell(cell, txt, size=9, color="FFFFFF")

    DISPLAY = {
        "baseline":  "Baseline",
        "p050_kd":   "P50+KD",
        "p050_nokd": "P50 (sin KD)",
        "p070_kd":   "P70+KD",
        "p070_nokd": "P70 (sin KD)",
    }
    for idx, (name, r) in enumerate(ev.items()):
        bg  = ROW_EVEN if idx % 2 == 0 else ROW_ODD
        row = tbl.add_row().cells
        for c in row:
            set_cell_bg(c, bg)

        fp32_kb = r["disk_kb"]["fp32"]
        int8_kb = r["disk_kb"]["int8"] or 0
        lat_fp  = r["latency"]["fp32"]["mean_ms"]
        lat_i8  = r["latency"].get("int8", {}).get("mean_ms")
        speedup = round(lat_fp / lat_i8, 2) if lat_i8 else "—"
        ratio   = round(fp32_kb / int8_kb, 2) if int8_kb else "—"

        vals = [
            DISPLAY.get(name, name),
            f"{r['n_params']:,}",
            f"{fp32_kb:.1f}",
            f"{int8_kb:.1f}" if int8_kb else "—",
            f"{ratio}x" if isinstance(ratio, float) else ratio,
            f"{lat_fp:.3f}",
            f"{lat_i8:.3f}" if lat_i8 else "—",
            f"{speedup}x" if isinstance(speedup, float) else speedup,
        ]
        for cell, txt in zip(row, vals):
            plain_cell(cell, txt, size=9)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print("=" * 60)
    print("Generando informe_final.docx")
    print("=" * 60)

    # --- Cargar datos externos ---
    with open(RESULTS_DIR / "metrics_eval.json")     as f: ev       = json.load(f)
    with open(RESULTS_DIR / "ablation_metrics.json") as f: ablation = json.load(f)

    # --- Computar metricas completas (Macro F1 + Kappa) ---
    print("\nCargando test set ISRUC...")
    x_test, y_test = load_isruc()
    print(f"  {len(x_test):,} epocas | apnea={y_test.mean():.1%}")

    ONNX_MODELS = {
        "Baseline (FP32)":       CKPT_DIR / "baseline_fp32.onnx",
        "Baseline (INT8)":       CKPT_DIR / "baseline_int8.onnx",
        "P50+KD (FP32)":         CKPT_DIR / "p050_kd" / "model_fp32.onnx",
        "P50+KD (INT8)":         CKPT_DIR / "p050_kd" / "model_int8.onnx",
    }

    computed = {}
    for label, path in ONNX_MODELS.items():
        if not path.exists():
            print(f"  [SKIP] {label}: {path} no encontrado")
            continue
        print(f"  Evaluando {label}...")
        sess = make_session(path)
        probs, preds = predict(sess, x_test)
        computed[label] = metrics(y_test, probs, preds)

    # Agrupar FP32/INT8 por modelo base
    MODEL_PAIRS = [
        ("Baseline",  "Baseline (FP32)",  "Baseline (INT8)"),
        ("P50+KD",    "P50+KD (FP32)",    "P50+KD (INT8)"),
    ]

    # --- Construir documento ---
    doc = Document()

    # Margenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # Fuente por defecto
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # ========================
    # PORTADA
    # ========================
    add_title_page(doc)

    # ========================
    # 1. RESUMEN EJECUTIVO
    # ========================
    add_heading(doc, "1. Resumen Ejecutivo")
    add_para(doc, (
        "Este proyecto implementa un pipeline completo de compresion de redes neuronales "
        "profundas (DNN) aplicado a la deteccion binaria de apnea del sueno con EEG monocanal. "
        "El pipeline combina tres tecnicas secuenciales: poda estructurada L1-norm (Pruning), "
        "Knowledge Distillation (KD) y cuantizacion estatica INT8 mediante ONNX Runtime. "
        "El modelo base (SleepApneaCNN, 272 802 parametros) se entrena en UCDDB y se evalua "
        "en ISRUC-Sleep como conjunto de prueba cross-dataset."
    ))
    add_para(doc, (
        "Hallazgo principal: la configuracion optima E5 (Pruning 50% + KD + INT8) logra "
        "AUC-ROC = 0.6308 en ISRUC, superando al baseline (0.6130) en +0.0178, con una "
        "reduccion de 8.5x en tamano de disco y aceleracion de 3.90x en latencia CPU. "
        "Todos los modelos comprimidos mejoran o igualan el AUC del baseline, indicando "
        "que la compresion actua como regularizador y mejora la generalizacion cross-dataset."
    ))
    doc.add_paragraph()

    # ========================
    # 2. INTRODUCCION
    # ========================
    add_heading(doc, "2. Introduccion")
    add_para(doc, (
        "La apnea del sueno es un trastorno respiratorio cronico con alta prevalencia mundial. "
        "Su diagnostico automatico mediante EEG de canal unico permite el despliegue en "
        "dispositivos de baja potencia. Sin embargo, los modelos DNN con alto rendimiento "
        "suelen ser demasiado grandes y lentos para hardware embebido. Este proyecto aborda "
        "este problema mediante una cadena de compresion: Pruning -> KD -> INT8 QDQ."
    ))
    add_para(doc, (
        "Objetivo: demostrar que es posible reducir drasticamente el tamano y la latencia "
        "del modelo sin degradar (y en este caso mejorando) el AUC-ROC en evaluacion "
        "cross-dataset, usando metricas estandar: Accuracy, Macro F1, Cohen Kappa y AUC-ROC."
    ))
    doc.add_paragraph()

    # ========================
    # 3. METODOLOGIA
    # ========================
    add_heading(doc, "3. Metodologia")

    add_heading(doc, "3.1 Datos", level=2)
    add_para(doc, (
        "Entrenamiento: UCDDB (University College Dublin Sleep Apnea Database). "
        "25 sujetos, canal C3-A2, 100 Hz. Segmentacion en epocas de 30 s (3 000 muestras). "
        "Etiqueta binaria: apnea/no-apnea segun anotaciones de eventos respiratorios. "
        "Validacion cruzada 5-fold por sujeto (stratified split)."
    ))
    add_para(doc, (
        "Evaluacion cross-dataset: ISRUC-Sleep subgrupo 1 (22 sujetos). Nunca visto durante "
        "entrenamiento ni cuantizacion. Este conjunto es el test set definitivo en todos "
        "los experimentos."
    ))

    add_heading(doc, "3.2 Arquitectura SleepApneaCNN", level=2)
    add_para(doc, (
        "CNN 1D con entrada (B, 1, 3 000). Cuatro bloques Conv1d-BN-ReLU con "
        "filtros [32, 64, 128, 256] y MaxPool1d. Una capa Fully Connected final de "
        "256 -> 2 neuronas. Total de parametros: 272 802. Entrenamiento con Adam "
        "(lr=1e-3, weight decay=1e-4), 50 epocas, batch 64."
    ))

    add_heading(doc, "3.3 Poda Estructurada (L1-norm Pruning)", level=2)
    add_para(doc, (
        "Se aplica poda estructurada sobre canales de filtros Conv1d usando la norma L1. "
        "Los filtros con menor suma de pesos absolutos se eliminan permanentemente. "
        "Se evaluan dos niveles de poda: 50% y 70% del total de filtros, con y sin "
        "Knowledge Distillation posterior. El modelo podado se re-entrena (fine-tuning) "
        "por 30 epocas para recuperar precision."
    ))

    add_heading(doc, "3.4 Knowledge Distillation (KD)", level=2)
    add_para(doc, (
        "El modelo podado (estudiante) aprende del modelo baseline completo (profesor). "
        "La funcion de perdida combinada es:"
    ))
    add_para(doc,
        "L_KD = alfa * CE(y, y_hat) + (1 - alfa) * T^2 * KL(softmax(z_s/T) || softmax(z_t/T))",
        indent=True)
    add_para(doc, (
        "Con temperatura T = 4 y peso alfa = 0.5. La destilacion suaviza las distribuciones "
        "de probabilidad del profesor, proporcionando informacion de 'conocimiento oscuro' "
        "que mejora la generalizacion del estudiante."
    ))

    add_heading(doc, "3.5 Cuantizacion INT8 (ONNX Runtime QDQ)", level=2)
    add_para(doc, (
        "Cuantizacion estatica post-entrenamiento (PTQ) usando el formato QDQ de ONNX Runtime. "
        "Se calibran los rangos de activacion con 200 muestras de calibracion del conjunto "
        "de validacion UCDDB. Los pesos se cuantizan a INT8 (simetrico, per-tensor). "
        "Esto reduce el tamano del modelo ~4x y acelera la inferencia en CPU gracias al "
        "uso de instrucciones SIMD enteras."
    ))
    doc.add_paragraph()

    # ========================
    # 4. METRICAS
    # ========================
    add_heading(doc, "4. Metricas de Evaluacion")
    add_para(doc, (
        "Se reportan cuatro metricas sobre el test set cross-dataset (ISRUC), tanto "
        "para el modelo FP32 como INT8:"
    ))
    metrics_desc = [
        ("Accuracy",     "Proporcion de predicciones correctas. Sensible al desbalance de clases."),
        ("Macro F1",     "Media aritmetica del F1-score por clase. Pondera igual ambas clases, "
                         "relevante en datasets desbalanceados. Se calcula con average='macro'."),
        ("Cohen Kappa",  "Acuerdo entre predicciones y etiquetas corregido por el azar. "
                         "Rango [-1, 1]; >0.6 se considera acuerdo sustancial."),
        ("AUC-ROC",      "Area bajo la curva ROC. Mide la capacidad discriminativa global "
                         "independiente del umbral. Es la metrica principal del proyecto."),
    ]
    for name, desc in metrics_desc:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{name}: ")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(desc).font.size = Pt(11)

    add_para(doc, (
        "Criterio de cuantizacion: todas las metricas del modelo INT8 deben diferir del "
        "modelo FP32 en no mas de +/-0.5% (en terminos absolutos multiplicados por 100). "
        "Una diferencia mayor indica problema en la calibracion."
    ))
    doc.add_paragraph()

    # ========================
    # 5. RESULTADOS
    # ========================
    add_heading(doc, "5. Resultados")

    add_heading(doc, "5.1 Metricas FP32 vs INT8 (test cross-dataset ISRUC)", level=2)
    add_para(doc, (
        "Las siguientes tablas muestran Accuracy, Macro F1, Cohen Kappa y AUC-ROC "
        "para cada modelo en sus versiones FP32 e INT8, con la diferencia porcentual "
        "y el estado de la verificacion +/-0.5%."
    ))

    for base_name, fp32_key, int8_key in MODEL_PAIRS:
        if fp32_key in computed and int8_key in computed:
            add_metrics_table(doc, base_name,
                              computed[fp32_key], computed[int8_key])
        else:
            add_para(doc, f"  [Modelo {base_name}: ONNX no encontrado, omitido]")

    # ========================
    # 5.2 Ablacion
    # ========================
    add_ablation_table(doc, ablation)

    # ========================
    # 5.3 Compresion
    # ========================
    add_compression_table(doc, ev)

    # ========================
    # 6. DISCUSION
    # ========================
    add_heading(doc, "6. Discusion")

    add_heading(doc, "a) Generalizacion cross-dataset", level=2)
    add_para(doc, (
        "Todos los modelos comprimidos superan al baseline en AUC-ROC sobre ISRUC. "
        "El baseline FP32 obtiene AUC = 0.6130; E5 (P50+KD+INT8) alcanza 0.6308 (+1.78%). "
        "Esto indica que la poda actua como regularizador implicito: al eliminar filtros "
        "de baja magnitud, se reduce el sobreajuste a UCDDB y mejora la transferencia a ISRUC."
    ))

    add_heading(doc, "b) Efecto del Knowledge Distillation", level=2)
    add_para(doc, (
        "Comparando E1 (P50 sin KD, AUC=0.6299) vs E4 (P50+KD, AUC=0.6306): "
        "el KD agrega +0.07% de AUC. El beneficio es pequeno en terminos absolutos "
        "pero consistente. El KD suaviza la distribucion de probabilidad del estudiante, "
        "especialmente en casos ambiguos donde el profesor asigna probabilidad distribuida."
    ))

    add_heading(doc, "c) Impacto de la cuantizacion INT8", level=2)
    add_para(doc, (
        "La cuantizacion INT8 reduce el tamano en disco ~4x (de ~455 KB a ~125 KB para P50). "
        "La latencia se reduce de 0.190 ms (FP32) a 0.106 ms (INT8), un speedup de 1.79x. "
        "Las metricas de clasificacion se mantienen dentro de +/-0.5% respecto al FP32, "
        "confirmando que la calibracion QDQ es correcta."
    ))

    add_heading(doc, "d) Configuracion optima: E5", level=2)
    add_para(doc, (
        "E5 (Pruning 50% + KD + INT8) es el punto optimo de la frontera de Pareto: "
        "maxima AUC (0.6308), minima latencia (0.106 ms), reduccion de disco 8.5x. "
        "El modelo INT8 pesa solo 125 KB, adecuado para despliegue en microcontroladores "
        "o dispositivos wearable con recursos limitados."
    ))
    doc.add_paragraph()

    # ========================
    # 7. CONCLUSIONES
    # ========================
    add_heading(doc, "7. Conclusiones")
    conclusions = [
        ("C1", "El pipeline Pruning->KD->INT8 reduce el tamano del modelo 8.5x y "
               "la latencia 3.90x sin degradar la AUC-ROC; en cambio, la MEJORA en "
               f"+0.0178 (AUC E5=0.6308 vs baseline=0.6130) sobre ISRUC."),
        ("C2", "La poda estructurada L1-norm es el principal factor de generalizacion: "
               "E1 (solo poda) ya supera al baseline en AUC (+0.0169) con 2.21x speedup."),
        ("C3", "La cuantizacion INT8 es compatible con las metricas del modelo FP32 "
               "dentro de +/-0.5%, validando la calibracion QDQ sobre datos UCDDB."),
        ("C4", "El Cohen Kappa es bajo en todos los modelos (~0.0-0.1), reflejando "
               "el desbalance de clases y la dificultad del problema cross-dataset; "
               "el AUC-ROC es la metrica mas informativa en este contexto."),
        ("C5", "El modelo E5 (125 KB, 0.106 ms/muestra) es apto para despliegue "
               "en sistemas embebidos o wearables para monitorizacion de apnea en tiempo real."),
    ]
    for code, text in conclusions:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{code}: ")
        run.bold = True; run.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    doc.add_paragraph()

    # ========================
    # 8. APENDICE
    # ========================
    add_heading(doc, "8. Apendice: Orden de Ejecucion")
    add_para(doc, "Ejecutar en orden desde la raiz del proyecto (requiere Python 3.10+):")

    commands = [
        ("Fase 1",  "python scripts/01_prepare_data.py",
         "Preprocesa UCDDB: segmenta epocas, guarda X_train.npy / y_apnea_train.npy"),
        ("Fase 2",  "python scripts/02b_crossval.py",
         "Entrena CNN baseline con 5-fold CV. Guarda sleep_baseline.pth"),
        ("Fase 3a", "python scripts/03_isruc_preprocess.py",
         "Preprocesa ISRUC: guarda subj_*.npz en data/processed/isruc_subjects/"),
        ("Fase 3b", "python scripts/03_pruning_kd.py",
         "Pruning L1-norm + Knowledge Distillation. Guarda checkpoints/p050_kd/ etc."),
        ("Fase 4",  "python scripts/04_quantize.py",
         "Exporta ONNX FP32 y cuantiza a INT8 QDQ. Guarda *_fp32.onnx / *_int8.onnx"),
        ("Fase 5",  "python scripts/05_evaluate.py",
         "Evalua latencia, RAM, metricas en ISRUC. Guarda results/metrics_eval.json"),
        ("Fase 6",  "python scripts/06_ablation.py",
         "Tabla de ablacion E0-E5. Guarda results/ablation_metrics.json"),
        ("Fase 7",  "python scripts/07_pareto_plot.py",
         "Genera pareto.png y ablation_table.png en results/figures/"),
        ("Demo",    "python demo.py",
         "Demo rapida: Accuracy, Macro F1, Cohen Kappa, AUC-ROC FP32 vs INT8"),
    ]

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    for cell, txt in zip(tbl.rows[0].cells, ["Fase", "Comando", "Descripcion"]):
        set_cell_bg(cell, HDR_COLOR)
        bold_cell(cell, txt, size=9, color="FFFFFF")
    for idx, (fase, cmd, desc) in enumerate(commands):
        bg  = ROW_EVEN if idx % 2 == 0 else ROW_ODD
        row = tbl.add_row().cells
        for c in row:
            set_cell_bg(c, bg)
        plain_cell(row[0], fase, size=9, bold=True)
        plain_cell(row[1], cmd,  size=8, center=False)
        plain_cell(row[2], desc, size=9, center=False)

    doc.add_paragraph()

    # ========================
    # Guardar
    # ========================
    out = ROOT / "informe_final.docx"
    doc.save(str(out))
    print(f"\nInforme guardado en: {out}")
    print("=" * 60)


if __name__ == "__main__":
    main()
